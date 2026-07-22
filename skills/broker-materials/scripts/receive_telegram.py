#!/usr/bin/env python3
"""receive_telegram.py — 券商材料接收器（B5a 完整實作）

用法:
    python3 receive_telegram.py --test-file /tmp/report.pdf  # 直接處理本地檔案
    python3 receive_telegram.py --once                        # 單次（供 OpenClaw 呼叫）

設計說明:
    Telegram getUpdates 已被 OpenClaw long-polling 佔用。
    正式流程: OpenClaw 收到 PDF → 存 /tmp → 呼叫本腳本 --test-file。
    本腳本不做 Telegram polling，只做 PDF 處理 + Notion 寫入 + Telegram 回報。
"""
import argparse
import json
import subprocess
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path

import requests
import re


# 股票代碼中英文對照（從 news_publisher.py 同步）
COMPANY_NAMES = {
    '2330': '台積電', '2454': '聯發科', '2303': '聯電', '2317': '鴻海',
    '2408': '南亞科', '2881': '富邦金', '2344': '華邦電', '2337': '旺宏',
    '8299': '群聯', '3037': '欣興', '4958': '臻鼎', '3712': '大聯大',
    '6175': '聯亞', '8086': '宏捷科', '6274': '台燿', '2383': '台光電',
    '3122': '萬潤', '3711': '日月光', '8261': '矽品', '3189': '景碩',
    '3105': '穩懋', '2409': '友達', '3481': '群創', '6515': '穎崴',
    '2345': '智邦', '2327': '國巨', '6116': '彩晶', '2357': '華碩',
    '2353': '宏碁', '2376': '技嘉', '2377': '微星', '6150': '撼訊',
    '3443': '創意', '3035': '智原', '3661': '世芯', '6531': '愛普',
    '3529': '力旺', '3131': '弘塑', '6939': '天虹', '2456': '全新',
    '6213': '聯茂', '6269': '台郡', '2313': '華通', '2367': '燿華',
    '8046': '南電', '3528': '景崎', '1560': '中砂', '6442': '光聖',
    '6285': '啟碁', '3312': '至上', '1319': '東陽',
}


WORKSPACE = Path.home() / ".openclaw" / "workspace"
SECRETS_FILE = WORKSPACE / "config" / "secrets.json"
PDF_READER = WORKSPACE / "skills" / "pdf-reader" / "scripts" / "pdf_dispatch.py"

VALID_CATEGORIES = ["stock_report", "industry_report", "morning_brief"]
CLASSIFY_TEXT_HEAD = 8000
CLASSIFY_TEXT_TAIL = 2000

MINIMAX_BASE = "https://api.minimax.io/anthropic/v1"
MINIMAX_MODEL = "MiniMax-M2.7"
NOTION_API = "https://api.notion.com/v1"
NOTION_VERSION = "2022-06-28"

LOG_FILE = Path("/tmp/receive_telegram.log")

def _write_log(msg: str):
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(msg + "\n")


def load_secrets():
    return json.loads(SECRETS_FILE.read_text(encoding="utf-8"))


def pdf_to_text(pdf_path: Path) -> str:
    # 純文字檔案直接讀取，不走 pdf-reader
    if pdf_path.suffix.lower() in (".txt", ".md"):
        return pdf_path.read_text(encoding="utf-8", errors="replace")
    
    # Word 檔（.docx / .doc）
    if pdf_path.suffix.lower() in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(str(pdf_path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if r: parts.append(r)
            t = "\n".join(parts)
            if not t.strip(): raise RuntimeError("Word 內容為空")
            return t
        except ImportError:
            raise RuntimeError("需安裝 python-docx: pip install python-docx --break-system-packages")
        except Exception as e:
            raise RuntimeError(f"Word 讀取失敗: {e}")

    # ZIP 壓縮檔
    if pdf_path.suffix.lower() == ".zip":
        import zipfile, tempfile; texts = []
        try:
            with zipfile.ZipFile(pdf_path, "r") as z:
                ms = [
                    n for n in z.namelist()
                    if not n.endswith("/")
                    and Path(n).suffix.lower() in (".pdf",".docx",".doc",".txt",".md")]
                if not ms: raise RuntimeError("ZIP 內無可讀檔案")
                for name in ms:
                    tp = Path(tempfile.mktemp(suffix=Path(name).suffix.lower()))
                    try:
                        with z.open(name) as fz: tp.write_bytes(fz.read())
                        ft = pdf_to_text(tp)
                        if ft.strip(): texts.append(f"=== {Path(name).name} ===\n{ft}")
                    except Exception as ie: texts.append(f"=== {Path(name).name} === [失敗:{ie}]")
                    finally:
                        if tp.exists(): tp.unlink()
        except zipfile.BadZipFile as e: raise RuntimeError(f"ZIP 損壞:{e}")
        if not texts: raise RuntimeError("ZIP 內所有檔案讀取失敗")
        return "\n\n".join(texts)

    # 嘗試 pdfplumber，若文字層內容豐富則直接使用
    script_dir = PDF_READER.parent.resolve()
    sys.path.insert(0, str(script_dir))
    try:
        # run pdfplumber in a child process: hard 90s kill on hang
        worker = Path(__file__).parent / "_pdfplumber_worker.py"
        r = subprocess.run(
            [sys.executable, str(worker), str(pdf_path), str(script_dir)],
            capture_output=True, text=True, timeout=90,
        )
        if r.returncode == 0:
            text = r.stdout
            import re
            cleaned = re.sub(r"--- Page \d+ ---", "", text)
            cleaned = re.sub(r"\s+", "", cleaned)
            if len(cleaned) >= 300:
                return text
    except Exception:
        pass
    
    # Step 2: 文字層不足，改走 OCR
    try:
        from extract_ocr import extract_with_ocr
        text, _ = extract_with_ocr(pdf_path)
        if len(text.strip()) > 0:
            return text
    except Exception:
        pass
    
    # Fallback: dispatch（最後手段）
    result = subprocess.run(
        [sys.executable, str(PDF_READER), str(pdf_path)],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"pdf-reader 失敗: {result.stderr}")
    return result.stdout



def _extract_stock_codes(text, limit=8):
    """4-digit TW-code candidates with year/number-noise filtering.
    - reject if adjacent to digits, comma, dot, slash, dash, or
      year/month/day markers (kills 6,852 / 4.49 / 2026年 / 10/07)
    - reject 2020..2040 as years unless whitelisted in COMPANY_NAMES
    """
    pat = (r"(?<![\d,./\-])([12][0-9]{3})"
           r"(?![\d,./\-年月日%])")
    out = []
    for c in dict.fromkeys(re.findall(pat, text)):
        if c in COMPANY_NAMES:
            out.append(c)
        elif 2020 <= int(c) <= 2040:
            continue
        else:
            out.append(c)
    return out[:limit]

def trim_for_classify(text: str) -> str:
    """取前 8000 + 後 2000 字元，避免超出 M2.7 甜蜜點。"""
    if len(text) <= CLASSIFY_TEXT_HEAD + CLASSIFY_TEXT_TAIL:
        return text
    return text[:CLASSIFY_TEXT_HEAD] + "\n...[中略]...\n" + text[-CLASSIFY_TEXT_TAIL:]


def classify_with_m27(text: str, secrets: dict) -> dict:
    """M2.7 分類 + 萃取結構化欄位。thinking=off（事實萃取類）。

    5層 fallback：
    1. 直接 json.loads(raw)（可能截斷）
    2. 截斷 JSON 修復
    3. ast.literal_eval（Python dict repr 格式）
    4. Markdown 表格解析（針對多 stock morning_brief）
    5. Thinking block 文字分析（針對只有 thinking 的回應）
    """
    import re, ast

    trimmed = trim_for_classify(text)

    system_prompt = (
        "你是台股券商報告分類器。收到報告後，直接輸出純 JSON，不加任何說明文字。\n\n"
        "分類規則：\n"
        "- stock_report：針對單一個股的研究報告（含評等/目標價）\n"
        "- industry_report：產業趨勢或主題報告\n"
        "- morning_brief：每日晨報（市場綜覽/多股簡評）\n\n"
 "特別規則（Call Memo / 法說會摘要）：\n"
 "- 若文件標題或開頭出現「Call Memo」字樣，代表這是法說會/電話會議摘要。\n"
 "- broker_name 必須填寫實際發布機構名稱（通常標示在「Call Memo」字樣之前，如「國泰證期研究部」），\n"
 " 絕對不可以把「Call Memo」「CALLMEMO」這幾個字本身當作 broker_name。\n"
 "- 若原文只有單一日期（通常在「Call Memo」下方，如 20260311），請將該日期格式化為 YYYY-MM-DD，\n"
 " 填入 investor_meeting_date，視為法說會/電話會議舉辦日期。\n"
 "- rating/target_price 若原文沒有明確評等或目標價（Call Memo 常見），請填 未明確/0。\n\n"
        'stock_report 格式（數值欄位填數字，無資料填0）：\n'
        '{"category":"stock_report","confidence":"high/medium/low","report_date":"YYYY-MM-DD",'
        '"stock_code":"如2330.TW","company_name":"公司名","broker_name":"券商名",'
        '"rating":"買進/加碼/中立/減碼/賣出/未明確","target_price":0,"current_price":0,'
        '"core_view":"核心觀點","revenue_forecast_this_year":0,"revenue_forecast_next_year":0,'
        '"eps_forecast_this_year":0,"eps_forecast_next_year":0,"gross_margin_forecast":0,'
        '"pe_valuation":0,"key_excerpt":"關鍵段落200字內","investor_meeting_date":"YYYY-MM-DD或空字串"}\n\n'
        'industry_report 格式：\n'
        '{"category":"industry_report","confidence":"high/medium/low","report_date":"YYYY-MM-DD",'
        '"topic":"主題標題","industry_tags":["產業"],"broker_name":"券商名","core_view":"核心觀點",'
        '"key_numbers":"關鍵數字","beneficiary_stocks":["代碼"],"risk_stocks":["代碼"],'
        '"key_excerpt":"關鍵段落200字內"}\n\n'
        'morning_brief 格式：\n'
        '{"category":"morning_brief","confidence":"high/medium/low","report_date":"YYYY-MM-DD",'
        '"broker_name":"券商名","core_view":"重點3句內"}'
    )

    headers = {
        "Content-Type": "application/json",
        "x-api-key": secrets["minimax_api_key"],
    }
    payload = {
        "model": MINIMAX_MODEL,
        "max_tokens": 6000,
        "thinking": {"type": "disabled"},
        "system": system_prompt,
        "messages": [
            {"role": "user", "content": f"請分類並萃取以下報告：\n\n{trimmed}"},
        ],
    }
    resp = requests.post(f"{MINIMAX_BASE}/messages", headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    blocks = resp.json()["content"]

    # Layer 0: collect text blocks (filter empty, prefer longest)
    candidates = [b["text"] for b in blocks if b.get("type") == "text" and b.get("text", "").strip()]
    text_blocks = sorted(candidates, key=len, reverse=True) if candidates else []

    if not text_blocks:
        # Layer 5: 只有 thinking block → 從 thinking 文字分析
        result = _fallback_from_thinking(blocks)
        if result:
            return result
        raise RuntimeError(f"M2.7 回應無有效 text block: {[b.get('type') for b in blocks]}")

    raw = text_blocks[0].strip().lstrip("\n")  # 移除前導換行
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1].rsplit("```", 1)[0].strip()

    # Layer 1: 直接 JSON（可能截斷）
    if raw.startswith("{"):
        try:
            return json.loads(raw)
        except json.JSONDecodeError as e:
            if e.msg in ("Unterminated string starting at", "Expecting value",
                         "Unterminated object", "Invalid control character"):
                result = _try_complete_json(raw)
                if result:
                    print("  [FALLBACK] truncated JSON 修復成功")
                    return result

    # Layer 2: 修補（找到第一個 {，並嘗試解析）
    original_raw = raw
    if not raw.startswith("{"):
        brace_pos = raw.find("{")
        if brace_pos > 0:
            raw = raw[brace_pos:]
        elif brace_pos == -1:
            pass  # 純 markdown，直接交給 Layer 4

    if raw.startswith("{"):
        try:
            return json.loads(raw)
        except json.JSONDecodeError as e:
            if e.msg in ("Unterminated string starting at", "Expecting value",
                         "Unterminated object", "Invalid control character"):
                result = _try_complete_json(raw)
                if result:
                    print("  [FALLBACK] truncated JSON 修復成功 (L2)")
                    return result

    # Layer 3: ast.literal_eval（Python dict repr）
    try:
        parsed = ast.literal_eval(raw)
        if isinstance(parsed, dict) and parsed.get("category"):
            print("  [FALLBACK] ast.literal_eval 成功")
            return parsed
    except (ValueError, SyntaxError):
        pass

    # Layer 4: Markdown 表格解析
    result = _fallback_from_markdown(original_raw)
    if result:
        return result

    # Last resort: 嘗試從 thinking block 擷取
    result = _fallback_from_thinking(blocks)
    if result:
        return result
    raise RuntimeError(f"M2.7 回應無法解析: {raw[:100]}")


def _try_complete_json(raw: str) -> dict | None:
    """嘗試修復被截斷的 JSON。"""
    import json

    # Method 1: 直接解析
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    # Method 2: 找最後一個完整閉合點截斷
    markers = ['"}', '\n]']
    for m in markers:
        idx = raw.rfind(m)
        if idx > 0:
            candidate = raw[:idx + len(m)]
            try:
                result = json.loads(candidate)
                if result.get("category"):
                    return result
            except:
                pass

    # Method 3: 砍掉最後一行再試
    lines = raw.split("\n")
    if len(lines) > 1:
        for i in range(len(lines)-1, max(0, len(lines)-10), -1):
            candidate = "\n".join(lines[:i])
            try:
                result = json.loads(candidate)
                if result.get("category"):
                    return result
            except:
                pass

    return None



def _fallback_from_markdown(md_raw: str) -> dict | None:
    """從 M2.7 回傳的 markdown 表格格式中萃取欄位。"""
    import re

    if md_raw.count("{") > 0 or "|" not in md_raw:
        return None

    fields = {}
    has_bold_key = False

    # First pass: **bold** key format: | **Key** | value |
    for line in md_raw.split("\n"):
        line = line.strip()
        if not line.startswith("|") or "---" in line or ":--" in line:
            continue
        parts = [p.strip() for p in line.split("|")]
        if len(parts) >= 3:
            key_cell = parts[1]
            val_cell = parts[2]
            km = re.search(r"\*\*([^*]+)\*\*", key_cell)
            if km:
                has_bold_key = True
                key = km.group(1).strip()
                val = re.sub(r"\*+", "", val_cell).strip().rstrip("/").strip()
                if val:
                    fields[key] = val

    # Second pass: header-row mapping (multi-stock table: | 代碼 | 公司 | 評等 |)
    if not has_bold_key and not fields:
        table_lines = [l.strip() for l in md_raw.split("\n")
                       if l.strip().startswith("|") and "---" not in l and ":--" not in l]
        if len(table_lines) >= 2:
            header_parts = [p.strip() for p in table_lines[0].split("|")]
            header_parts = [p for p in header_parts if p]
            for data_line in table_lines[1:]:
                data_parts = [p.strip() for p in data_line.split("|")]
                data_parts = [p for p in data_parts if p]
                for col_idx in range(min(len(header_parts), len(data_parts))):
                    header = header_parts[col_idx]
                    value = re.sub(r"\*+", "", data_parts[col_idx]).strip().rstrip("/").strip()
                    if value and header:
                        fields[header] = value

    if not fields:
        return None

    # Multi-stock detection: stock code like "6239 TT" in column 1
    stock_code_col = fields.get("股票代碼", "") or fields.get("代碼", "") or ""
    company_col = fields.get("公司名稱", "") or fields.get("公司", "") or ""
    rating_col = fields.get("投資評等", "") or fields.get("評等", "") or ""
    target_col = fields.get("目標價", "") or fields.get("目標價位", "") or ""

    is_stock_code_value = bool(re.match(r"^\d{4}\s*[Tt][Tt]$", stock_code_col))

    if is_stock_code_value and rating_col:
        result = {
            "category": "morning_brief",
            "confidence": "medium",
            "broker_name": fields.get("報告來源", "") or fields.get("券商", "") or "其他",
            "core_view": f"股票：{stock_code_col} {company_col}，評等：{rating_col}，目標價：{target_col}"
        }
        print(f"  [FALLBACK] markdown multi-stock: {stock_code_col}")
        return result

    # Single-stock detection
    rtype = fields.get("報告類型", "") or fields.get("文件類型", "") or ""
    if any(k in rtype for k in ["晨報", "盤後", "晨訊", "晨會"]):
        result = {"category": "morning_brief", "confidence": "low",
                  "broker_name": fields.get("報告來源", "") or fields.get("券商", "") or "其他",
                  "core_view": str(fields)}
        print(f"  [FALLBACK] markdown: morning_brief")
        return result
    elif "產業報告" in rtype:
        result = {"category": "industry_report", "confidence": "low",
                  "topic": fields.get("研究產業", "") or "未命名",
                  "broker_name": fields.get("報告來源", "") or "其他",
                  "core_view": str(fields)}
        print(f"  [FALLBACK] markdown: industry_report")
        return result
    else:
        stock = fields.get("目標公司", "") or fields.get("公司名稱", "") or ""
        result = {"category": "stock_report", "confidence": "low",
                  "stock_code": fields.get("股票代碼", ""), "company_name": stock,
                  "broker_name": fields.get("報告來源", "") or "其他",
                  "rating": fields.get("評等", "未明確"),
                  "core_view": str(fields), "key_excerpt": str(fields)}
        print(f"  [FALLBACK] markdown: stock_report")
        return result


def _extract_broker_name(thinking: str) -> str:
    """從 thinking 文字萃取券商名稱，找不到回傳 '未知'。"""
    for kw in ["富邦","元大","國泰","凱基","統一","永豐金","群益","兆豐",
                "中信","日盛","華南","台新","第一金","玉山","遠東","福邦",
                "大和","野村","麥格理","美林","摩根","高盛","瑞銀","德意志",
                "匯豐","花旗","法巴","巴克萊"]:
        if kw in thinking: return kw
    return "未知"


def _fallback_from_thinking(blocks: list) -> dict | None:
    """當 M2.7 只回傳 thinking block 時，從文字內容萃取欄位。"""
    import re

    for b in blocks:
        if b.get("type") == "thinking" and b.get("thinking", "").strip():
            thinking = b.get("thinking", "")

            # Find all stock codes: "6239 TT" pattern
            all_codes = re.findall(r"(\d{4})\s+[Tt][Tt]", thinking)

            rating_m = re.search(r"(強力買進|買進|持有|中立|減碼|賣出)", thinking)
            target_m = re.search(r"目標[價價位]?\s*(\d+(?:\.\d+)?)", thinking)
            is_multi = any(k in thinking for k in ["3家", "三家", "涵蓋", "包含", "多家", "多檔"])
            is_brief = any(k in thinking for k in ["晨報", "晨訊", "盤後", "投顧股市"])

            if is_brief or is_multi or len(all_codes) > 1:
                result = {
                    "category": "morning_brief",
                    "confidence": "low",
                    "broker_name": _extract_broker_name(thinking),
                    "core_view": f"來源：{_extract_broker_name(thinking)}報告。涵蓋股票：{', '.join(all_codes) if all_codes else '未識別'}. {thinking[:300]}"
                }
                print(f"  [FALLBACK] thinking: morning_brief (codes={all_codes})")
                return result
            elif all_codes:
                stock_code = all_codes[0] + ".TW"
                rating = rating_m.group(1) if rating_m else "未明確"
                tp = target_m.group(1) if (target_m and target_m.group(1).replace(".", "").isdigit()) else "0"
                result = {
                    "category": "stock_report",
                    "confidence": "low",
                    "stock_code": stock_code,
                    "company_name": "",
                    "broker_name": _extract_broker_name(thinking),
                    "rating": rating,
                    "target_price": float(tp) if tp != "0" else 0,
                    "core_view": f"來源：{_extract_broker_name(thinking)}報告。評等：{rating}。目標價：{tp}元。內容：{thinking[:300]}",
                    "key_excerpt": thinking[:500]
                }
                print(f"  [FALLBACK] thinking: stock_report ({stock_code})")
                return result

    return None




def rt(content: str) -> list:
    """Notion rich_text 格式，截斷 2000 字元。"""
    return [{"text": {"content": str(content)[:2000]}}]


def write_investor_meeting(stock_code: str, meeting_date: str, secrets: dict):
    """若券商報告提到法說會日期，寫入 event_calendar（idempotent）。"""
    import time as _time
    db_id = secrets["notion_event_calendar_db"]
    headers = {
        "Authorization": f"Bearer {secrets['notion_key']}",
        "Notion-Version": NOTION_VERSION,
        "Content-Type": "application/json",
    }
    # idempotency check
    check = requests.post(
        f"{NOTION_API}/databases/{db_id}/query",
        headers=headers,
        json={"filter": {"and": [
            {"property": "股票代碼", "rich_text": {"equals": stock_code}},
            {"property": "事件類型", "select": {"equals": "法說會"}},
            {"property": "預計日期", "title": {"equals": meeting_date}},
        ]}, "page_size": 1},
        timeout=15,
    )
    check.raise_for_status()
    if check.json().get("results"):
        print(f"  [SKIP] 法說會已存在: {stock_code} {meeting_date}")
        return
    props = {
        "預計日期": {"title": [{"text": {"content": meeting_date}}]},
        "股票代碼": {"rich_text": [{"text": {"content": stock_code}}]},
        "事件類型": {"select": {"name": "法說會"}},
        "重要性":   {"select": {"name": "高"}},
        "已提醒":   {"checkbox": False},
    }
    resp = requests.post(
        f"{NOTION_API}/pages",
        headers=headers,
        json={"parent": {"database_id": db_id}, "properties": props},
        timeout=15,
    )
    resp.raise_for_status()


def _make_content_blocks(core_view: str, key_excerpt: str) -> list:
    """生成 Notion 頁面內文 blocks（核心觀點 + 關鍵摘錄）。"""
    blocks = [
        {"object":"block","type":"heading_2","heading_2":{"rich_text":[{"text":{"content":"📊 核心觀點"}}]}},
        {"object":"block","type":"paragraph","paragraph":{"rich_text":[{"text":{"content":core_view[:2000]}}]}},
        {"object":"block","type":"divider","divider":{}},
    ]
    if key_excerpt.strip():
        blocks += [
            {"object":"block","type":"heading_2","heading_2":{"rich_text":[{"text":{"content":"📄 關鍵摘錄"}}]}},
            {"object":"block","type":"paragraph","paragraph":{"rich_text":[{"text":{"content":key_excerpt[:2000]}}]}},
        ]
    return blocks


_INVALID_BROKER_MARKERS = ["CALLMEMO", "MEMO", "展會", "觀察", "筆記"]

def _sanitize_broker_name(name) -> str:
    name = str(name or "").strip()
    if not name:
        return "其他"
    upper = name.upper()
    if any(m in upper or m in name for m in _INVALID_BROKER_MARKERS):
        return "其他"
    return name


# ==== Call Memo full-text archive (added 2026-07-21) ====
# NOTE: all Chinese literals below are \uXXXX-escaped so this whole block is
# pure ASCII end-to-end (transfer-channel safety, see pitfalls #59/#62/#64).

RE_CM_LINE = re.compile(r"^\s*[Cc][Aa][Ll][Ll]\s*[Mm][Ee][Mm][Oo]\s*$")
RE_CM_DATE = re.compile(r"^\s*(20\d{6})\s*$")
RE_CM_COMPANY = re.compile("^(.{1,20}?)[\uff08(](\\d{4})(?:\\s*TT)?[\uff09)]\\s*$")
RE_CM_BROKER = re.compile("(\u7814\u7a76\u90e8|\u8b49\u5238|\u6295\u9867|\u6295\u4fe1|Securities)")
RE_CM_SIGNAL = re.compile("(callmemo|(?:\u6cd5\u8aaa|\u6cd5\u4eba\u8aaa\u660e\u6703|\u6cd5\u4eba\u5ea7\u8ac7\u6703).{0,8}(?:\u7d00\u8981|\u6458\u8981|memo))")
RE_CM_ANYDATE_SLASH = re.compile(r"(?<!\d)(20\d{2})[/.\-](\d{1,2})[/.\-](\d{1,2})(?!\d)")
RE_CM_ANYDATE_8 = re.compile(r"(?<!\d)(20\d{2})(0[1-9]|1[0-2])([0-2]\d|3[01])(?!\d)")
RE_CM_DATE6 = re.compile(r"(?<!\d)(2[4-9])(0[1-9]|1[0-2])([0-2]\d|3[01])(?!\d)")
RE_CM_ANYCOMPANY = re.compile("(?<![\u4e00-\u9fff])([\u4e00-\u9fff][\u4e00-\u9fffA-Za-z0-9]{0,9}?)[ \\t]*[\uff08(](\\d{4})(?:\\s*TT)?[\uff09)]")
RE_CM_CMCODE = re.compile("[Cc]all\\s*[Mm]emo[_\\s]*(\\d{4})\\s*([\u4e00-\u9fff]{2,10})")
RE_CM_ANYBROKER = re.compile("([\u4e00-\u9fff]{2,8}?(?:\u8b49\u671f\u7814\u7a76\u90e8|\u7814\u7a76\u90e8|\u6295\u9867|\u8b49\u5238|\u6295\u4fe1))")
RE_CM_MEMOLINE = re.compile(r"^\s*[Mm][Ee][Mm][Oo]\s*$")
RE_CM_CODEFIRST = re.compile("[\uff08(](\\d{4})[\uff09)][ \\t]*([\u4e00-\u9fff][\u4e00-\u9fffA-Za-z0-9\\-]{1,11})")
CM_FNAME_KEYWORDS = ("callmemo", "\u6cd5\u8aaa")


def _cm_norm(s):
    return re.sub(r"[\s_\-]+", "", str(s or "")).lower()


def _cm_head_lines(text, n):
    return [l.strip() for l in text.split("\n") if l.strip()][:n]


def _cm_extract_date(head_text):
    m = RE_CM_ANYDATE_SLASH.search(head_text)
    if m:
        mo = int(m.group(2))
        d = int(m.group(3))
        if 1 <= mo <= 12 and 1 <= d <= 31:
            return "%s-%02d-%02d" % (m.group(1), mo, d)
    m = RE_CM_ANYDATE_8.search(head_text)
    if m:
        return m.group(1) + "-" + m.group(2) + "-" + m.group(3)
    m = RE_CM_DATE6.search(head_text)
    if m:
        return "20" + m.group(1) + "-" + m.group(2) + "-" + m.group(3)
    return ""


def _cm_extract_company(head_text):
    for m in RE_CM_ANYCOMPANY.finditer(head_text):
        name = m.group(1).strip()
        code = m.group(2)
        if 2020 <= int(code) <= 2040 and code not in COMPANY_NAMES:
            continue
        return name, code
    for m in RE_CM_CODEFIRST.finditer(head_text):
        code = m.group(1)
        name = m.group(2).strip()
        if 2020 <= int(code) <= 2040 and code not in COMPANY_NAMES:
            continue
        return name, code
    m = RE_CM_CMCODE.search(head_text)
    if m:
        return m.group(2).strip(), m.group(1)
    return "", ""


def _cm_extract_broker(head_text):
    m = RE_CM_ANYBROKER.search(head_text)
    return m.group(1) if m else ""


CM_EN_BROKERS = (("kgi", "\u51f1\u57fa\u6295\u9867"),
                 ("yuanta", "\u5143\u5927\u6295\u9867"),
                 ("cathay", "\u570b\u6cf0\u8b49\u671f\u7814\u7a76\u90e8"),
                 ("ctbc", "\u4e2d\u4fe1\u6295\u9867"),
                 ("fubon", "\u5bcc\u90a6\u6295\u9867"),
                 ("sinopac", "\u6c38\u8c50\u91d1\u8b49\u5238"))
CM_CATHAY_MARK = "\u50c5\u4f9b\u672c\u516c\u53f8\u5167\u90e8\u540c\u4ec1\u53c3\u8003\u4f7f\u7528\uff0c\u975e\u7d93\u672c\u516c\u53f8\u4e8b\u5148\u66f8\u9762\u540c\u610f"


def _cm_broker_fallback(seg):
    if seg["broker"]:
        return seg
    head = _cm_norm("\n".join(_cm_head_lines(seg["text"], 8)))
    for en, zh in CM_EN_BROKERS:
        if en in head:
            seg["broker"] = zh
            return seg
    if CM_CATHAY_MARK in seg["text"]:
        seg["broker"] = "\u570b\u6cf0\u8b49\u671f\u7814\u7a76\u90e8"
    return seg


def split_call_memos(text):
    """Tier-1 deterministic split on standalone 'Call Memo' line + 8-digit date.
    Returns list of {text, date, company, code, broker}; [] when no anchor."""
    lines = text.split("\n")
    anchors = []
    for i, l in enumerate(lines):
        if RE_CM_LINE.match(l):
            date = None
            for j in range(i + 1, min(i + 4, len(lines))):
                m = RE_CM_DATE.match(lines[j])
                if m:
                    date = m.group(1)
                    break
                if lines[j].strip():
                    break
            if date:
                anchors.append((i, date))
    if not anchors:
        return []
    starts = []
    for i, _d in anchors:
        start = i
        k = i - 1
        steps = 0
        while k >= 0 and steps < 4:
            ls = lines[k].strip()
            if ls and (RE_CM_COMPANY.match(ls) or RE_CM_BROKER.search(ls)):
                start = k
                k -= 1
                steps += 1
            elif not ls:
                k -= 1
                steps += 1
            else:
                break
        starts.append(start)
    segs = []
    for n, (i, date) in enumerate(anchors):
        s = starts[n]
        e = starts[n + 1] if n + 1 < len(anchors) else len(lines)
        seg_text = "\n".join(lines[s:e]).strip("\n")
        broker = ""
        comp = ""
        code = ""
        for l in lines[s:i]:
            ls = l.strip()
            m = RE_CM_COMPANY.match(ls)
            if m:
                comp, code = m.group(1).strip(), m.group(2)
            elif RE_CM_BROKER.search(ls) and not broker:
                broker = ls
        iso = date[:4] + "-" + date[4:6] + "-" + date[6:]
        segs.append({"text": seg_text, "date": iso, "company": comp,
                     "code": code, "broker": broker})
    return segs


def detect_call_memos(text, file_name):
    """Returns (segments, mode). mode: 'anchor' | 'fallback' | ''.
    anchor   = Cathay-style standalone 'Call Memo' + 8-digit date header.
    fallback = filename or per-line head-zone call-memo signal; whole text
               as one segment, metadata via generic header extractors
               (KGI/CTBC/Yuanta style headers)."""
    segs = split_call_memos(text)
    if segs:
        return [_cm_broker_fallback(s) for s in segs], "anchor"
    fn = _cm_norm(file_name)
    hit = any(k in fn for k in CM_FNAME_KEYWORDS)
    if not hit:
        for l in _cm_head_lines(text, 15):
            if RE_CM_SIGNAL.search(_cm_norm(l)):
                hit = True
                break
    if not hit:
        for l in text.split("\n")[:6]:
            if RE_CM_MEMOLINE.match(l):
                hit = True
                break
    if not hit:
        return [], ""
    head8 = "\n".join(_cm_head_lines(text, 8))
    name, code = _cm_extract_company(head8)
    return [_cm_broker_fallback({"text": text.strip("\n"),
             "date": _cm_extract_date(head8),
             "company": name,
             "code": code,
             "broker": _cm_extract_broker(head8)})], "fallback"


def chunk_full_text_blocks(text, limit=1900):
    """Split full text into Notion paragraph blocks along paragraph/newline
    boundaries; hard-cut only when a single run exceeds the limit."""
    paras = text.split("\n\n")
    pieces = []
    buf = ""
    for p in paras:
        cand = (buf + "\n\n" + p) if buf else p
        if len(cand) <= limit:
            buf = cand
            continue
        if buf:
            pieces.append(buf)
            buf = ""
        while len(p) > limit:
            cut = p.rfind("\n", 0, limit)
            if cut < limit // 2:
                cut = limit
            pieces.append(p[:cut])
            p = p[cut:].lstrip("\n")
        buf = p
    if buf:
        pieces.append(buf)
    return [{"object": "block", "type": "paragraph",
             "paragraph": {"rich_text": [{"text": {"content": pc}}]}}
            for pc in pieces if pc.strip()]


def archive_call_memos(text, file_name, m27, secrets, source="\u7ba1\u7dda", audit=None):
    """Additive full-text archive into the dedicated call-memo Notion DB.
    Dedupe key: md5 of segment text (content-md5 property). Per-segment
    failures only WARN; returns number of pages created."""
    import hashlib
    import time as _t
    db_id = secrets.get("notion_callmemo_db")
    if not db_id:
        print("  [MEMO] notion_callmemo_db not in secrets, skip archive")
        return 0
    segs, mode = detect_call_memos(text, file_name)
    if not segs:
        return 0
    m27 = m27 or {}
    headers = {
        "Authorization": "Bearer " + secrets["notion_key"],
        "Notion-Version": NOTION_VERSION,
        "Content-Type": "application/json",
    }
    today = datetime.now(timezone(timedelta(hours=8))).strftime("%Y-%m-%d")
    created = 0
    lines_ok = []
    for seg in segs:
        rec = {"file": str(file_name), "mode": mode}
        try:
            seg_text = seg["text"]
            md5 = hashlib.md5(seg_text.encode("utf-8")).hexdigest()
            rec.update({"md5": md5, "chars": len(seg_text)})
            q = requests.post(
                NOTION_API + "/databases/" + db_id + "/query", headers=headers,
                json={"filter": {"property": "\u5167\u6587MD5",
                                 "rich_text": {"equals": md5}},
                      "page_size": 1},
                timeout=15)
            q.raise_for_status()
            if q.json().get("results"):
                print("  [MEMO][SKIP] duplicate md5=" + md5[:8] + " " + str(file_name))
                rec["action"] = "skip_dup"
                if audit is not None:
                    audit.append(rec)
                continue
            date = seg["date"]
            if not re.match(r"^\d{4}-\d{2}-\d{2}$", date or ""):
                date = ""
            if not date:
                for k in ("investor_meeting_date", "report_date"):
                    v = str(m27.get(k) or "").strip()
                    if re.match(r"^\d{4}-\d{2}-\d{2}$", v):
                        date = v
                        break
            if not date:
                date = today
            code = seg["code"] or re.sub(r"\D", "", str(m27.get("stock_code") or ""))[:4]
            name = seg["company"] or str(m27.get("company_name") or "").strip()
            broker = _sanitize_broker_name(seg["broker"] or m27.get("broker_name"))
            if name or code:
                title = name + "(" + code + ") " + date + " \u6cd5\u8aaa\u6703Memo"
            else:
                title = str(file_name) + " " + date + " \u6cd5\u8aaa\u6703Memo"
            props = {
                "\u6a19\u984c": {"title": [{"text": {"content": title[:100]}}]},
                "\u516c\u53f8\u4ee3\u78bc": {"rich_text": rt(code + ".TW" if code else "")},
                "\u516c\u53f8\u540d\u7a31": {"rich_text": rt(name)},
                "\u6cd5\u8aaa\u6703\u65e5\u671f": {"date": {"start": date}},
                "\u4f86\u6e90\u5238\u5546": {"select": {"name": broker[:100]}},
                "\u539f\u59cb\u6a94\u540d": {"rich_text": rt(str(file_name))},
                "\u5168\u6587\u5b57\u6578": {"number": len(seg_text)},
                "\u5167\u6587MD5": {"rich_text": rt(md5)},
                "\u4f86\u6e90": {"select": {"name": source}},
            }
            blocks = [{"object": "block", "type": "heading_2", "heading_2": {
                "rich_text": [{"text": {"content": "\U0001f4de \u6cd5\u8aaa\u6703\u5168\u6587"}}]}}]
            blocks += chunk_full_text_blocks(seg_text)
            first = blocks[:95]
            rest = blocks[95:]
            resp = requests.post(NOTION_API + "/pages", headers=headers,
                                 json={"parent": {"database_id": db_id},
                                       "properties": props, "children": first},
                                 timeout=30)
            resp.raise_for_status()
            page = resp.json()
            page_id = page.get("id", "")
            while rest:
                batch = rest[:95]
                rest = rest[95:]
                _t.sleep(0.4)
                ar = requests.patch(NOTION_API + "/blocks/" + page_id + "/children",
                                    headers=headers, json={"children": batch},
                                    timeout=30)
                ar.raise_for_status()
            created += 1
            rec.update({"action": "created", "date": date, "code": code,
                        "url": page.get("url", "")})
            if audit is not None:
                audit.append(rec)
            lines_ok.append((code or name or "?") + " " + date)
            print("  [MEMO][OK] " + title[:60] + " chars=" + str(len(seg_text)))
        except Exception as e:
            print("  [WARN] memo segment archive failed: " + str(e))
            rec.update({"action": "error", "error": str(e)[:200]})
            if audit is not None:
                audit.append(rec)
    if created and source == "\u7ba1\u7dda":
        try:
            notify_telegram("\U0001f4de \u6cd5\u8aaa\u6703Memo\u5df2\u5b8c\u6574\u6b78\u6a94 " + str(created) + " \u7bc7\uff5c"
                            + str(file_name) + "\n" + "\u3001".join(lines_ok[:6]),
                            secrets)
        except Exception:
            pass
    return created



def write_to_notion(category: str, fields: dict, secrets: dict) -> str:
    """寫入對應 Notion DB，回傳 page URL。morning_brief 不寫 Notion，回傳空字串。"""
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    _rd = str(fields.get("report_date") or "")
    report_date = _rd if re.match(r"^\d{4}-\d{2}-\d{2}$", _rd) else today
    try:
        _delta = abs((datetime.strptime(report_date, "%Y-%m-%d")
                      - datetime.strptime(today, "%Y-%m-%d")).days)
    except ValueError:
        _delta = 9999
    if _delta > 90:
        fields["core_view"] = ("[日期待確認:" + report_date + "] "
                               + str(fields.get("core_view", "")))
        report_date = today
    headers = {
        "Authorization": f"Bearer {secrets['notion_key']}",
        "Notion-Version": NOTION_VERSION,
        "Content-Type": "application/json",
    }

    if category == "stock_report":
        db_id = secrets["notion_broker_reports_db"]
        props = {
            "股票代碼": {"title": rt(fields.get("stock_code", "未知"))},
            "公司名稱": {"rich_text": rt(fields.get("company_name", ""))},
            "券商名稱": {"select": {"name": _sanitize_broker_name(fields.get("broker_name"))[:100]}},
            "評等": {"select": {"name": str(fields.get("rating") or "未明確")}},
            "報告日期": {"date": {"start": report_date}},
            "核心觀點": {"rich_text": rt(fields.get("core_view", ""))},
            "原始內文": {"rich_text": rt(fields.get("key_excerpt", ""))},
        }
        for py_key, notion_key in [
            ("target_price", "目標價"),
            ("current_price", "報告當日股價"),
            ("revenue_forecast_this_year", "營收預測_今年"),
            ("revenue_forecast_next_year", "營收預測_明年"),
            ("eps_forecast_this_year", "EPS預測_今年"),
            ("eps_forecast_next_year", "EPS預測_明年"),
            ("gross_margin_forecast", "毛利率預測"),
            ("pe_valuation", "PE估值"),
        ]:
            val = fields.get(py_key, 0)
            try:
                if val and float(val) != 0:
                    fval = float(val)
                    if py_key == "gross_margin_forecast" and abs(fval) > 1.5:
                        # Notion percent field stores 0.45 for 45%
                        fval = fval / 100.0
                    props[notion_key] = {"number": fval}
            except (TypeError, ValueError):
                pass

        # report-type marking: sparse fields => snippet, else full report
        _sparse = (
            "目標價" not in props
            and "報告當日股價" not in props
            and str(fields.get("rating") or "未明確") in ("未明確", "未評等")
        )
        props["報告型態"] = {"select": {"name": "短評/快訊" if _sparse else "完整報告"}}

    elif category == "industry_report":
        db_id = secrets["notion_industry_reports_db"]
        props = {
            "產業主題": {"title": [{"text": {"content": fields.get("topic", "未命名產業報告")}}]},
            "券商名稱": {"select": {"name": _sanitize_broker_name(fields.get("broker_name"))[:100]}},
            "報告日期": {"date": {"start": report_date}},
            "核心觀點": {"rich_text": rt(fields.get("core_view", ""))},
            "關鍵數字": {"rich_text": rt(fields.get("key_numbers", ""))},
            "原始內文_關鍵段落": {"rich_text": rt(fields.get("key_excerpt", ""))},
        }
        tags = fields.get("industry_tags") or []
        if tags:
            props["產業分類"] = {"multi_select": [{"name": str(t)[:100]} for t in tags[:5]]}
        for fkey, nkey in [("beneficiary_stocks", "受惠標的"), ("risk_stocks", "受害標的")]:
            stocks = fields.get(fkey) or []
            if stocks:
                props[nkey] = {"multi_select": [{"name": str(s)[:100]} for s in stocks[:10]]}

    elif category == "morning_brief":
        db_id = secrets["notion_industry_reports_db"]
        broker = str(fields.get("broker_name") or "未知")
        props = {
            "產業主題": {"title": [{"text": {"content": f"晨報｜{broker}"}}]},
            "券商名稱": {"select": {"name": broker[:100]}},
            "報告日期": {"date": {"start": report_date}},
            "核心觀點": {"rich_text": rt(fields.get("core_view", ""))},
            "關鍵數字": {"rich_text": rt("")},
            "原始內文_關鍵段落": {"rich_text": rt(fields.get("core_view", ""))},
            "產業分類": {"multi_select": [{"name": "晨報"}]},
        }
        stocks = fields.get("beneficiary_stocks") or []
        if stocks:
            props["受惠標的"] = {"multi_select": [{"name": str(s)[:100]} for s in stocks[:10]]}
        else:
            return ""
        # fall through to write


    props["digest_mark"] = {"rich_text": [{"text": {"content": "processed"}}]}
    children = _make_content_blocks(str(fields.get("core_view","")), str(fields.get("key_excerpt","")))
    body = {"parent": {"database_id": db_id}, "properties": props, "children": children}
    resp = requests.post(f"{NOTION_API}/pages", headers=headers, json=body, timeout=30)
    resp.raise_for_status()
    return resp.json().get("url", "")


def notify_telegram(message: str, secrets: dict):
    """POST Telegram Bot API sendMessage 給 Kai DM。"""
    url = f"https://api.telegram.org/bot{secrets['telegram_bot_token']}/sendMessage"
    try:
        resp = requests.post(url, json={
            "chat_id": secrets["telegram_dm"],
            "text": message,
            "parse_mode": "HTML",
        }, timeout=15)
        resp.raise_for_status()
    except Exception as e:
        print(f"[WARN] Telegram 通知失敗: {e}", file=sys.stderr)




def send_report_to_group(file_path: Path, caption: str, secrets: dict) -> bool:
    """Send original broker file to broker group (secrets['telegram_group']) for reconcile+archive.
    Failures only WARN; never break the Notion write."""
    group = secrets.get("telegram_group")
    if not group:
        print(" [WARN] no telegram_group in secrets, skip group send", file=sys.stderr)
        return False
    url = "https://api.telegram.org/bot" + secrets["telegram_bot_token"] + "/sendDocument"
    try:
        with open(file_path, "rb") as _fh:
            resp = requests.post(
                url,
                data={"chat_id": group, "caption": str(caption)[:1024]},
                files={"document": (file_path.name, _fh)},
                timeout=60,
            )
        if resp.status_code == 200 and resp.json().get("ok"):
            print(" [OK] file sent to broker group: " + file_path.name)
            return True
        print(" [WARN] group send failed: %s %s" % (resp.status_code, resp.text[:200]), file=sys.stderr)
        return False
    except Exception as e:
        print(" [WARN] group send error: %s" % e, file=sys.stderr)
        return False


def classify_with_m27_with_retry(text: str, secrets: dict) -> dict:
    """classify_with_m27 包裝，最多重試 3 次（Timeout/ConnError/5xx）。"""
    import time as _t
    _attempts, _delay = 3, 20
    for _i in range(_attempts):
        try:
            return classify_with_m27(text, secrets)
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as _e:
            if _i < _attempts - 1:
                print(f" [RETRY] M2.7 連線失敗，{_delay}s 後重試 ({_i+1}/{_attempts-1})...")
                _t.sleep(_delay)
            else:
                raise
        except requests.exceptions.HTTPError as _e:
            if _e.response.status_code >= 500 and _i < _attempts - 1:
                print(f" [RETRY] M2.7 伺服器錯誤 ({_e.response.status_code})，{_delay}s 後重試...")
                _t.sleep(_delay)
            else:
                raise

def process_file(file_path: Path, secrets: dict):
    line = f"[START] {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} {file_path.name}"; print(line); _write_log(line)

    # Step 1: PDF → text
    print("[1/4] pdf-reader 轉文字...")
    text = pdf_to_text(file_path)
    print(f"  → {len(text)} 字元")

    # Step 2: M2.7 分類 + 萃取
    print("[2/4] M2.7 分類中...")
    result = classify_with_m27_with_retry(text, secrets)
    category = result.get("category", "unknown")
    confidence = result.get("confidence", "low")
    print(f"  → category={category}, confidence={confidence}")

    # Step 3: confidence / category 分流
    if category not in VALID_CATEGORIES:
        msg = (f"⚠️ 無法分類\n檔案: {file_path.name}\n"
               f"M2.7 回傳: {json.dumps(result, ensure_ascii=False)[:300]}")
        notify_telegram(msg, secrets)
        print(f"[WARN] 分類失敗，已通知 Kai")
        return

    if confidence == "low":
        msg = (f"🤔 分類信心低，請人工確認\n"
               f"檔案: {file_path.name}\n"
               f"分類: {category}\n"
               f"摘要: {json.dumps(result, ensure_ascii=False)[:400]}")
        notify_telegram(msg, secrets)
        result["core_view"] = "[待確認] " + str(result.get("core_view", ""))
        print("[INFO] 信心低，標記 [待確認] 後繼續寫 Notion")

    # morning_brief: 注入股票代碼清單供 write_to_notion 使用
    if category == "morning_brief":
        result["beneficiary_stocks"] = _extract_stock_codes(text)

    # Step 4: 寫入 Notion
    print("[3/4] 寫入 Notion...")
    # Step 3b: Call Memo full-text archive (additive, before main write)
    try:
        _n_memo = archive_call_memos(text, file_path.name, result, secrets)
        if _n_memo:
            print("  [MEMO] archived " + str(_n_memo) + " segment(s)")
    except Exception as _me:
        print("  [WARN] call memo archive failed: " + str(_me))

    page_url = write_to_notion(category, result, secrets)

    # Step 5: Telegram 回報
    print("[4/4] Telegram 通知...")
    if category == "stock_report":
        msg = (f"✅ 券商個股報告已存入\n"
               f"📌 {result.get('stock_code','')} {result.get('company_name','')}\n"
               f"🏦 {result.get('broker_name','')} | 評等: {result.get('rating','')}\n"
               f"🎯 目標價: {result.get('target_price','')} | 現價: {result.get('current_price','')}\n"
               f"💬 {str(result.get('core_view',''))[:120]}\n"
               f"🔗 {page_url}")
    elif category == "industry_report":
        tags_str = ", ".join(result.get("industry_tags") or [])
        msg = (f"✅ 券商產業報告已存入\n"
               f"📋 {result.get('topic','')}\n"
               f"🏦 {result.get('broker_name','')} | {tags_str}\n"
               f"💬 {str(result.get('core_view',''))[:120]}\n"
               f"🔗 {page_url}")
    else:  # morning_brief
        # 存入今日晨報合併檔
        # 使用台北時間（UTC+8），與 broker_digest.py 的 date_compact 一致
        TPE = timezone(timedelta(hours=8))
        today_str = datetime.now(TPE).strftime("%Y%m%d")
        morning_file = WORKSPACE / "state" / f"broker_morning_{today_str}.txt"
        morning_file.parent.mkdir(parents=True, exist_ok=True)
        broker_name = result.get('broker_name', '')
        report_date = result.get('report_date', '')
        core_view = str(result.get('core_view', ''))
        # 取出股票代碼清單
        found_codes = _extract_stock_codes(text)
        code_parts = []
        for c in found_codes:
            name = COMPANY_NAMES.get(c, '')
            code_parts.append(f"{c} {name}" if name else c)
        code_line = "、".join(code_parts) if code_parts else "（未偵測到代碼）"
        with open(morning_file, "a", encoding="utf-8") as f:
            f.write(f"\n{'='*40}\n")
            f.write(f"券商：{broker_name}\n")
            f.write(f"日期：{report_date}\n")
            f.write(f"股票：{code_line}\n")
            f.write(f"核心觀點：{core_view}\n")
            f.write(f"原文：{text[:3000]}\n")
        # 一行確認
        notify_telegram(f"✅ 已收晨報 {broker_name}｜股票：{code_line}\n🔗 {page_url}", secrets)

    # Step 5c: send original file to broker group (telegram_group) for reconcile + archive
    if category == 'stock_report':
        _cap = (str(result.get('stock_code','')) + ' ' + str(result.get('company_name','')) + ' | '
            + str(result.get('broker_name','')) + ' | ' + str(result.get('rating','')) + ' | TP '
            + str(result.get('target_price','')) + '\n' + str(page_url))
    elif category == 'industry_report':
        _cap = (str(result.get('topic','')) + ' | ' + str(result.get('broker_name','')) + '\n' + str(page_url))
    else:
        _cap = (str(result.get('broker_name','')) + ' morning brief\n' + str(page_url))
    send_report_to_group(file_path, _cap, secrets)

 # Step 5b: 若 stock_report 有法說會日期，寫入 event_calendar
    if category == "stock_report":
        meeting_date = str(result.get("investor_meeting_date", "")).strip()
        stock_code = result.get("stock_code", "")
        if meeting_date and len(meeting_date) == 10 and meeting_date != "YYYY-MM-DD":
            try:
                write_investor_meeting(stock_code, meeting_date, secrets)
                print(f"  [OK] 法說會事件寫入 event_calendar: {stock_code} {meeting_date}")
            except Exception as e:
                print(f"  [WARN] 法說會寫入失敗: {e}")

    line = f"[DONE] {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} {category}"; print(line); _write_log(line)


def main():
    parser = argparse.ArgumentParser(description="Telegram 券商材料接收")
    parser.add_argument("--once", action="store_true", help="單次處理（OpenClaw 整合用）")
    parser.add_argument("--test-file", help="直接測試本地 PDF 檔案")
    parser.add_argument("--text", help="直接處理本地純文字檔案（.txt/.md）")
    args = parser.parse_args()

    secrets = load_secrets()

    if args.test_file:
        file_path = Path(args.test_file).expanduser().resolve()
        if not file_path.exists():
            print(f"[ERROR] 檔案不存在: {file_path}", file=sys.stderr)
            return 1
        process_file(file_path, secrets)
        return 0

    if args.text:
        file_path = Path(args.text).expanduser().resolve()
        if not file_path.exists():
            print(f"[ERROR] 檔案不存在: {file_path}", file=sys.stderr)
            return 1
        process_file(file_path, secrets)
        return 0

    # getUpdates 被 OpenClaw long-polling 佔用，本腳本不做 polling
    print("[INFO] Telegram polling 由 OpenClaw 負責。")
    print("[INFO] 使用方式: python3 receive_telegram.py --test-file /tmp/report.pdf")
    return 0


if __name__ == "__main__":
    sys.exit(main())